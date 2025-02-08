# Standard library imports
import os

# third-party imports
import fitz
from docx import Document
from pathlib import Path
import uuid
import logging
from PIL import Image

# local imports
from utils.config import collection, text_model, FIXED_DIMENSION, text_splitter, image_model
from utils.helpers import adjust_embedding_dimension
from src.model import run_caption_model

logger = logging.getLogger(__name__)

def process_text(file_path):
    """
    Process a text file, generate embeddings, and store in ChromaDB.

    Args:
        file_path (str): Path to the text file.
    """
    print(f"Processing text file: {file_path}")
    with open(file_path, "r", encoding="utf-8") as f:
        text_content = f.read()

    # Split text into chunks
    chunks = text_splitter.split_text(text_content)
    
    if not chunks:
        print(f"No content found in file: {file_path}")
        return False

    documents_list = []
    embeddings_list = []
    ids_list = []
    metadatas_list = []

    for i, chunk in enumerate(chunks):
        # Generate embedding for each chunk
        embedding = text_model.encode(chunk)
        adjusted_embedding = adjust_embedding_dimension(embedding, FIXED_DIMENSION)
        
        # Create unique ID
        chunk_id = f"text_{Path(file_path).stem}_{uuid.uuid4()}"
        
        # Prepare metadata
        metadata = {
            "type": "text",
            "filename": Path(file_path).name,
            "file_path": str(file_path),
            "file_type": "text",
            "chunk_index": i,
            "total_chunks": len(chunks),
            "content_length": len(chunk),
            "creation_date": str(os.path.getctime(file_path))
        }

        documents_list.append(chunk)
        embeddings_list.append(adjusted_embedding)
        ids_list.append(chunk_id)
        metadatas_list.append(metadata)

    # Only add to collection if we have content
    if documents_list and embeddings_list:
        collection.add(
            embeddings=embeddings_list,
            documents=documents_list,
            ids=ids_list,
            metadatas=metadatas_list,
        )
        print(f"Text file '{file_path}' processed and added to ChromaDB.")
        return True
    else:
        print(f"No valid content to add from file: {file_path}")
        return False


def process_pdf(file_path):
    """
    Process a PDF file, extract text, generate embeddings, and store in ChromaDB.
    
    Args:
        file_path (str): Path to the PDF file
    
    Returns:
        bool: True if successful, False otherwise
    """
    try:
        print(f"Processing PDF file: {file_path}")
        
        # Validate file exists
        if not Path(file_path).exists():
            raise FileNotFoundError(f"PDF file not found: {file_path}")
            
        # Extract content from PDF
        try:
            doc = fitz.open(file_path)
            text_list = []
            
            # Get PDF metadata
            page_count = len(doc)
            pdf_title = doc.metadata.get("title", "")
            pdf_author = doc.metadata.get("author", "")
            pdf_creation_date = doc.metadata.get("creationDate", "")
            
            # Process each page
            for page_num, page in enumerate(doc):
                # Extract text
                text = page.get_text()
                if text.strip():
                    text_list.append(text)
                    
            if not text_list:
                raise ValueError("PDF appears to be empty or contains no extractable content")
                
            # Split text into chunks
            chunks = text_splitter.split_text("\n".join(text_list)) if text_list else []
                
        except Exception as e:
            raise ValueError(f"Failed to extract content from PDF: {str(e)}")
        finally:
            doc.close()
            
        # Process and store content
        try:
            # Process text chunks
            if chunks:
                text_documents = []
                text_embeddings = []
                text_ids = []
                text_metadatas = []
                
                for i, chunk in enumerate(chunks):
                    embedding = text_model.encode(chunk)
                    adjusted_embedding = adjust_embedding_dimension(embedding, FIXED_DIMENSION)
                    
                    chunk_id = f"pdf_text_{Path(file_path).stem}_{uuid.uuid4()}"
                    chunk_metadata = {
                        "type": "text",
                        "filename": Path(file_path).name,
                        "file_path": str(file_path),
                        "file_type": "pdf",
                        "chunk_index": i,
                        "total_chunks": len(chunks),
                        "page_count": str(page_count),
                        "pdf_title": pdf_title,
                        "pdf_author": pdf_author,
                        "pdf_creation_date": pdf_creation_date
                    }
                    
                    text_documents.append(chunk)
                    text_embeddings.append(adjusted_embedding)
                    text_ids.append(chunk_id)
                    text_metadatas.append(chunk_metadata)
                
                # Store text chunks
                collection.add(
                    embeddings=text_embeddings,
                    documents=text_documents,
                    ids=text_ids,
                    metadatas=text_metadatas,
                )
            
            print(f"PDF file '{file_path}' processed and added to ChromaDB.")
            return True
                
        except Exception as e:
            raise RuntimeError(f"Failed to add content to ChromaDB: {str(e)}")
            
    except Exception as e:
        logger.error(f"Error processing PDF {file_path}: {str(e)}")
        print(f"Failed to process PDF {file_path}: {str(e)}")
        return False

def process_docx(file_path):
    """
    Process a DOCX file, extract text, generate embeddings, and store in ChromaDB.
    
    Args:
        file_path (str): Path to the DOCX file
    
    Returns:
        bool: True if successful, False otherwise
    """
    try:
        print(f"Processing DOCX file: {file_path}")
        
        # Validate file exists
        if not Path(file_path).exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
        
        # Load and extract text from DOCX
        try:
            doc = Document(file_path)
            full_text = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():  # Skip empty paragraphs
                    full_text.append(para.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text.strip())
            
            document_text = "\n".join(full_text)
            
            if not document_text.strip():
                raise ValueError("Document appears to be empty")
            
        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
        
        # Generate text embedding
        try:
            embedding = text_model.encode(document_text)
            adjusted_embedding = adjust_embedding_dimension(embedding, FIXED_DIMENSION)
        except Exception as e:
            raise RuntimeError(f"Failed to generate text embedding: {str(e)}")
        
        # Add to ChromaDB
        try:
            collection.add(
                documents=[document_text],
                embeddings=[adjusted_embedding],
                metadatas=[{
                    "type": "text",
                    "filename": Path(file_path).name,
                    "file_path": str(file_path),
                    "file_type": "docx",
                    "content_length": len(document_text),
                    "paragraph_count": len(doc.paragraphs),
                    "table_count": len(doc.tables),
                    "creation_date": str(os.path.getctime(file_path))
                }],
                ids=[f"docx_{uuid.uuid4()}"]
            )
            print(f"DOCX file '{file_path}' processed and added to ChromaDB.")
            return True
            
        except Exception as e:
            raise RuntimeError(f"Failed to add document to ChromaDB: {str(e)}")
            
    except Exception as e:
        logger.error(f"Error processing DOCX {file_path}: {str(e)}")
        print(f"Failed to process DOCX {file_path}: {str(e)}")
        return False

